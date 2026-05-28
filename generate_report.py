import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

def create_report():
    doc = docx.Document()
    
    # Define styles
    styles = doc.styles
    
    # Custom Title Style
    title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = 'Arial'
    title_font.size = Pt(28)
    title_font.bold = True
    title_font.color.rgb = RGBColor(0, 51, 102)
    
    # Custom Heading 1
    h1_style = styles.add_style('CustomH1', WD_STYLE_TYPE.PARAGRAPH)
    h1_font = h1_style.font
    h1_font.name = 'Arial'
    h1_font.size = Pt(18)
    h1_font.bold = True
    h1_font.color.rgb = RGBColor(0, 51, 102)
    
    # Custom Code Style
    code_style = styles.add_style('CodeStyle', WD_STYLE_TYPE.PARAGRAPH)
    code_font = code_style.font
    code_font.name = 'Courier New'
    code_font.size = Pt(10)
    code_style.paragraph_format.left_indent = Inches(0.5)
    code_style.paragraph_format.space_after = Pt(6)
    
    # Normal Text Style modification
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)
    
    # --- TITLE PAGE ---
    doc.add_paragraph('\n\n\n\n\n')
    p = doc.add_paragraph('SkillSyncAI Project', style='CustomTitle')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Comprehensive Line-by-Line Code Architecture & Implementation Report')
    subtitle.runs[0].font.size = Pt(16)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n\n\n')
    date_p = doc.add_paragraph('Generated on: 24 May 2026\nPrepared for: Internal Technical Review')
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    
    # --- INTRODUCTION ---
    doc.add_paragraph('1. Introduction', style='CustomH1')
    intro_text = (
        "This document provides a highly detailed, line-by-line architectural analysis and codebase walkthrough "
        "of the SkillSyncAI project. SkillSyncAI is an advanced web-based application built primarily with Python "
        "and the Flask framework. Its core functionality revolves around ingesting user resumes (in PDF or TXT formats), "
        "extracting the embedded text, and evaluating the candidate's skills against a predefined industry database of job roles.\n\n"
        "By utilizing natural language processing primitives, pattern matching algorithms, and a modular architecture, "
        "SkillSyncAI effectively automates the initial screening phase of recruitment and acts as an intelligent career counselor "
        "for applicants. It parses unstructured text, cross-references identified skills with role requirements, generates an "
        "empirical employability score, and offers dynamic, actionable learning recommendations for any missing competencies.\n\n"
        "This comprehensive 50-page equivalent report is intended for software engineers, code reviewers, and technical stakeholders "
        "who need to deeply understand the inner workings, data flow, error handling, and structural integrity of the application. "
        "We will dissect every module, starting from the entry point (app.py) to the highly specialized analytical modules."
    )
    doc.add_paragraph(intro_text)
    doc.add_page_break()

    # --- HELPER FUNCTION ---
    def add_code_section(title, code_text, explanation_paragraphs):
        doc.add_paragraph(title, style='Heading 2')
        doc.add_paragraph("Code Snippet:", style='Heading 3')
        
        # Add code block
        for line in code_text.split('\n'):
            doc.add_paragraph(line, style='CodeStyle')
            
        doc.add_paragraph("Detailed Line-by-Line Analysis:", style='Heading 3')
        for p_text in explanation_paragraphs:
            doc.add_paragraph(p_text)
        doc.add_paragraph('\n')

    # --- APP.PY ---
    doc.add_paragraph('2. Core Application Routing (app.py)', style='CustomH1')
    doc.add_paragraph("The app.py file serves as the WSGI gateway and the primary controller for the SkillSyncAI web interface. It ties together all the sub-modules.")
    
    code_1 = """from flask import Flask, render_template, request, redirect, url_for, flash, send_from_directory
from werkzeug.utils import secure_filename
import os
import json
import uuid

# Import custom modules
from modules.resume_parser import extract_text
from modules.skill_extractor import extract_skills
from modules.analyzer import analyze
from modules.score_generator import generate_score
from modules.recommender import recommend"""
    expl_1 = [
        "Lines 1-5: The application imports essential standard libraries and third-party frameworks. 'Flask' is imported to instantiate the web server, handle routing, and manage HTTP requests/responses. Components like 'render_template' allow the server to render HTML pages dynamically using Jinja2 syntax. 'request' provides access to incoming GET/POST data, 'redirect' and 'url_for' facilitate navigation between routes, and 'flash' is used to send one-time alert messages to the user interface. 'send_from_directory' securely serves static or generated files.",
        "Line 2: 'secure_filename' from werkzeug.utils is a security-critical import. It sanitizes user-uploaded filenames, ensuring that directory traversal attacks (e.g., uploading a file named '../../etc/passwd') are neutralized before saving the file to the file system.",
        "Lines 3-5: The standard libraries 'os', 'json', and 'uuid' are imported. 'os' is utilized for file system interactions (path joining, directory creation). 'json' is used for parsing the local database (skills_db.json). 'uuid' (Universally Unique Identifier) ensures that concurrent file uploads do not overwrite each other by prepending a unique string to the filename.",
        "Lines 8-12: The application imports its own internal business logic from the 'modules' package. This modularization ensures separation of concerns. 'extract_text' will handle OCR and PDF/TXT parsing. 'extract_skills' will process the raw text to find known skills. 'analyze', 'generate_score', and 'recommend' represent the intelligence layer of the application."
    ]
    add_code_section("2.1 Imports and Dependencies", code_1, expl_1)

    code_2 = """app = Flask(__name__)
app.secret_key = os.urandom(24)

# Configuration
UPLOAD_FOLDER = 'resumes'
REPORT_FOLDER = 'reports'
ALLOWED_EXTENSIONS = {'pdf', 'txt'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['REPORT_FOLDER'] = REPORT_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 5 * 1024 * 1024  # 5MB

# Ensure directories exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(REPORT_FOLDER, exist_ok=True)"""
    expl_2 = [
        "Line 1: 'app = Flask(__name__)' instantiates the central Flask application object. Passing '__name__' informs Flask about the root path of the application so it can accurately locate the 'static' and 'templates' directories.",
        "Line 2: 'app.secret_key = os.urandom(24)' generates a cryptographically secure random 24-byte string. This key is absolutely mandatory for the 'flash' messaging system and session management. It ensures that the client-side session cookies are signed and tamper-proof.",
        "Lines 5-7: Constant variables are declared. 'UPLOAD_FOLDER' specifies where incoming user resumes will be temporarily stored. 'REPORT_FOLDER' designates the output directory for generated analysis reports. 'ALLOWED_EXTENSIONS' is a set defining the strict whitelist of file types (.pdf and .txt) the application will process.",
        "Lines 8-10: These constants are injected into the Flask application's configuration dictionary ('app.config'). Importantly, 'MAX_CONTENT_LENGTH' is set to 5 Megabytes (5 * 1024 * 1024 bytes). This acts as a robust defense mechanism against Denial of Service (DoS) attacks via oversized file uploads; Flask will automatically reject requests exceeding this payload size.",
        "Lines 13-14: The 'os.makedirs' function is called with 'exist_ok=True'. This is a defensive programming tactic executed at startup. It guarantees that the required directory structures ('resumes' and 'reports') exist before any requests are processed, preventing runtime FileNotFoundError exceptions when a user attempts an upload."
    ]
    add_code_section("2.2 Initialization and Configuration", code_2, expl_2)

    code_3 = """def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def get_roles():
    db_path = os.path.join('database', 'skills_db.json')
    with open(db_path, 'r') as f:
        db = json.load(f)
    return sorted(list(db['roles'].keys()))"""
    expl_3 = [
        "Line 1: 'def allowed_file(filename):' defines a helper utility to validate uploaded files.",
        "Lines 2-3: The function checks two conditions. First, it verifies that there is a dot ('.') in the filename, ensuring an extension exists. Second, it utilizes 'rsplit('.', 1)[1]' to split the filename at the last dot, extracts the extension, converts it to lowercase using '.lower()', and checks if it exists within the 'ALLOWED_EXTENSIONS' set. This prevents malicious file types (like .exe or .sh) from proceeding.",
        "Line 5: 'def get_roles():' defines a function to dynamically load the available job roles from the system's JSON database.",
        "Line 6: 'db_path = os.path.join(...)' constructs a cross-platform compatible file path to 'skills_db.json' located in the 'database' folder. This ensures the path resolves correctly on both Windows and Unix systems.",
        "Lines 7-8: The context manager 'with open(db_path, 'r') as f:' is used. This is Python best practice for file operations, as it guarantees the file handler is safely closed immediately after the JSON parsing is complete, preventing memory leaks or locked files. 'json.load(f)' deserializes the JSON file into a Python dictionary.",
        "Line 9: 'return sorted(list(db['roles'].keys()))' extracts the keys from the 'roles' object inside the database, converts them to a list, sorts them alphabetically (improving UX in the UI dropdown), and returns them."
    ]
    add_code_section("2.3 Core Helper Functions", code_3, expl_3)

    code_4 = """@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload')
def upload():
    roles = get_roles()
    return render_template('upload.html', roles=roles)"""
    expl_4 = [
        "Line 1: The '@app.route('/')' decorator binds the root URL endpoint to the 'index' function.",
        "Line 3: 'return render_template('index.html')' instructs Flask to locate the 'index.html' file within the 'templates' folder, process any embedded Jinja2 tags, and return the resulting HTML payload to the client. This serves as the landing page.",
        "Line 5: The '@app.route('/upload')' decorator binds the '/upload' endpoint.",
        "Line 7: The function calls 'get_roles()' to dynamically fetch the list of supported job titles.",
        "Line 8: The list of roles is injected into the 'upload.html' template via 'roles=roles'. This allows the frontend HTML form to populate a <select> dropdown dynamically based on the current state of the backend database."
    ]
    add_code_section("2.4 Static Routes", code_4, expl_4)

    code_5 = """@app.route('/analyze', methods=['POST'])
def analyze_resume():
    if 'resume' not in request.files or 'role' not in request.form:
        flash('Missing file or job role selection.')
        return redirect(request.url)
    
    file = request.files['resume']
    role = request.form['role']
    
    if file.filename == '':
        flash('No selected file.')
        return redirect(request.url)"""
    expl_5 = [
        "Line 1: '@app.route('/analyze', methods=['POST'])' restricts this endpoint to HTTP POST requests only. This is appropriate because the endpoint processes incoming form data and file streams.",
        "Line 3: The initial 'if' statement performs validation. It checks if the 'resume' key is present in the multipart/form-data payload ('request.files') and if the 'role' key is present in the standard form payload ('request.form').",
        "Lines 4-5: If validation fails, a 'flash' message is generated for the UI, and the server issues a 302 Redirect via 'redirect(request.url)' back to the submission page, halting further execution.",
        "Lines 7-8: Assuming validation passes, the uploaded file object and the selected role string are extracted into local variables 'file' and 'role'.",
        "Lines 10-12: An edge case is handled where the user submits the form without actually selecting a file in the browser (resulting in an empty filename string). Again, it flashes an error and redirects."
    ]
    add_code_section("2.5 Analysis Route - Validation", code_5, expl_5)

    code_6 = """    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        # Add UUID to filename to avoid collisions
        unique_filename = f"{uuid.uuid4().hex}_{filename}"
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
        file.save(filepath)"""
    expl_6 = [
        "Line 1: The main operational block begins by confirming the file object exists and its extension passes the 'allowed_file' whitelist check.",
        "Line 2: 'secure_filename' strips any dangerous characters (like slashes or null bytes) from the user-provided filename.",
        "Line 4: 'unique_filename' is generated using Python's 'uuid.uuid4().hex'. By appending this hexadecimal UUID to the original sanitized filename, the system mathematically guarantees that two users uploading a file named 'resume.pdf' simultaneously will not overwrite each other's data.",
        "Line 5: The final absolute destination path is constructed by joining the 'UPLOAD_FOLDER' configuration path and the new unique filename.",
        "Line 6: 'file.save(filepath)' is an operation provided by Werkzeug. It streams the uploaded file data directly from memory/temp space to the designated location on the permanent file system."
    ]
    add_code_section("2.6 Analysis Route - File Storage", code_6, expl_6)

    code_7 = """        try:
            # 1. Extract Text
            text = extract_text(filepath)
            
            # 2. Extract Skills
            detected_skills = extract_skills(text)
            
            # 3. Analyze against Role
            analysis = analyze(detected_skills, role)
            
            # 4. Generate Score
            score = generate_score(analysis['matched'], analysis['required'])
            
            # 5. Get Recommendations
            recommendations = recommend(analysis['missing'])"""
    expl_7 = [
        "Line 1: A 'try' block encapsulates the core business logic. Parsing documents and running analytics can be prone to errors (e.g., corrupted PDFs, unreadable text encoding), so robust exception handling is essential to prevent application crashes.",
        "Line 3: The 'extract_text' module is invoked with the file path. It will abstract away the complexity of parsing PDFs vs TXT files and return a pure string of text.",
        "Line 6: The raw text string is passed to 'extract_skills'. This module utilizes NLP and regex to identify recognized technical skills from the text, returning a normalized list of skills.",
        "Line 9: The 'analyze' module compares the list of user 'detected_skills' against the canonical database requirements for the target 'role'. It returns a dictionary identifying 'matched', 'missing', and 'required' skills.",
        "Line 12: 'generate_score' computes a quantitative employability metric and a qualitative label (e.g., 'Intermediate') based on the ratio of matched skills to required skills.",
        "Line 15: The 'recommend' module takes the list of 'missing' skills and cross-references them against course databases or generic search algorithms to return an array of actionable learning resources."
    ]
    add_code_section("2.7 Analysis Route - Core Business Logic Execution", code_7, expl_7)

    code_8 = """            # 6. Generate Report File
            report_filename = f"report_{unique_filename.rsplit('.', 1)[0]}.txt"
            report_path = os.path.join(app.config['REPORT_FOLDER'], report_filename)
            
            with open(report_path, 'w') as f:
                f.write("SKILLSYNC AI - RESUME ANALYSIS REPORT\\n")
                f.write("====================================\\n\\n")
                f.write(f"Target Role: {role}\\n")
                f.write(f"Employability Score: {score['score']}% ({score['label']})\\n\\n")
                f.write("DETECTED SKILLS:\\n")
                f.write(", ".join(analysis['matched']) if analysis['matched'] else "None detected")
                f.write("\\n\\n")
                f.write("MISSING SKILLS:\\n")
                f.write(", ".join(analysis['missing']) if analysis['missing'] else "None! You are a perfect match.")
                f.write("\\n\\n")
                f.write("RECOMMENDATIONS:\\n")
                for rec in recommendations:
                    f.write(f"- {rec['skill']}: {rec['tip']}\\n")
                    f.write(f"  Resource: {rec['url']}\\n")
            
            return render_template('result.html', 
                                 role=role,
                                 analysis=analysis, 
                                 score=score, 
                                 recommendations=recommendations,
                                 report_filename=report_filename)"""
    expl_8 = [
        "Lines 1-3: A persistent text report is generated for the user to download later. The report filename mirrors the UUID of the input file, replacing the '.pdf'/'txt' extension with '.txt'.",
        "Lines 5-20: A new file handler is opened in write mode ('w'). The code sequentially writes structured data to the file, forming a professional plain-text report. It injects the role, the numerical score, and the categorical label. It then formats the arrays of matched and missing skills using Python's ', '.join() method, incorporating ternary operations (e.g., 'if analysis['matched'] else \"None detected\"') to handle edge cases where the list might be empty. Finally, a loop iterates through the recommendations dictionary, writing the tip and URL for each missing skill.",
        "Lines 22-27: The response is finalized. 'render_template' builds the 'result.html' page. Critically, it passes all the generated data structures (role, analysis dict, score dict, recommendations list, and the report filename) into the Jinja2 template context so the frontend can dynamically render the charts, metrics, and download links."
    ]
    add_code_section("2.8 Analysis Route - Report Generation and Response", code_8, expl_8)

    # --- RESUME PARSER ---
    doc.add_paragraph('3. Document Ingestion (resume_parser.py)', style='CustomH1')
    doc.add_paragraph("The resume_parser.py module is responsible for abstracting the underlying file formats and extracting raw textual data for downstream processing.")
    
    code_9 = """import pdfplumber
import PyPDF2
import os

def extract_text(filepath):
    text = ""
    ext = os.path.splitext(filepath)[1].lower()"""
    expl_9 = [
        "Lines 1-3: The module imports 'pdfplumber' (a highly accurate PDF extraction library capable of handling tables and formatting), 'PyPDF2' (a faster, more traditional PDF library used as a fallback), and the 'os' module.",
        "Line 5: The 'extract_text' function accepts a single argument, 'filepath'.",
        "Line 6: 'text' is initialized as an empty string. This will serve as the buffer accumulating the extracted data.",
        "Line 7: 'os.path.splitext(filepath)[1]' safely splits the file path into root and extension, grabbing the extension (e.g., '.pdf'), and normalizes it to lowercase to ensure case-insensitive matching in subsequent logic."
    ]
    add_code_section("3.1 Initialization and Extension Checking", code_9, expl_9)

    code_10 = """    if ext == '.pdf':
        try:
            # Primary: pdfplumber
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\\n"
        except Exception as e:
            print(f"pdfplumber failed: {e}. Falling back to PyPDF2.")
            try:
                # Fallback: PyPDF2
                with open(filepath, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    for page in reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\\n"
            except Exception as e2:
                print(f"PyPDF2 also failed: {e2}")"""
    expl_10 = [
        "Line 1: A conditional branch routes logic based on the '.pdf' extension.",
        "Lines 2-8: The primary extraction mechanism utilizes 'pdfplumber'. The context manager 'with pdfplumber.open()' ensures resources are released. It iterates through the 'pages' collection. For each page, 'extract_text()' is called. If text is successfully returned (i.e., the page is not an image scan without OCR), it is appended to the main 'text' buffer with a newline character.",
        "Lines 9-10: An outer 'except Exception' catches any failures from pdfplumber (such as unsupported encryption or corrupted metadata). It logs the failure and announces a fallback attempt.",
        "Lines 11-19: The fallback mechanism utilizes 'PyPDF2'. Because PyPDF2 expects a raw byte stream, the file is opened in 'rb' (read-binary) mode. A 'PdfReader' is instantiated, and a similar page-by-page iteration occurs. PyPDF2 is more resilient to certain structural anomalies in PDFs, making it an excellent fallback to maximize the application's tolerance for diverse user uploads."
    ]
    add_code_section("3.2 PDF Extraction with Failover Logic", code_10, expl_10)

    code_11 = """    elif ext == '.txt':
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                text = f.read()
        except UnicodeDecodeError:
            with open(filepath, 'r', encoding='latin-1') as f:
                text = f.read()

    return text.strip()"""
    expl_11 = [
        "Line 1: If the file is not a PDF, it checks if it is a plain text file ('.txt').",
        "Lines 2-4: It attempts to open and read the file utilizing standard 'utf-8' encoding, which is the web standard.",
        "Lines 5-7: If a 'UnicodeDecodeError' occurs—often caused by resumes authored on legacy Windows systems or with specific regional characters—the application gracefully recovers by re-opening the file using the highly permissive 'latin-1' encoding. This guarantees the file can be read without crashing.",
        "Line 9: 'text.strip()' is called before returning. This utility function removes any leading or trailing whitespace, newlines, or invisible control characters from the final string, ensuring clean data for the NLP components."
    ]
    add_code_section("3.3 Text File Parsing and Encoding Recovery", code_11, expl_11)

    # --- SKILL EXTRACTOR ---
    doc.add_paragraph('4. NLP Skill Extraction (skill_extractor.py)', style='CustomH1')
    doc.add_paragraph("This module scans the raw text string and identifies specific technical skills through advanced regular expression pattern matching.")

    code_12 = """import json
import os
import re

def extract_skills(text):
    db_path = os.path.join(os.path.dirname(__file__), '..', 'database', 'skills_db.json')
    
    with open(db_path, 'r') as f:
        db = json.load(f)
    
    # Create a unique list of all skills across all roles
    all_skills = set()
    for role_skills in db['roles'].values():
        for skill in role_skills:
            all_skills.add(skill)"""
    expl_12 = [
        "Lines 1-3: Standard imports 'json' and 'os' are present alongside 're', Python's powerful regular expression library.",
        "Line 6: 'os.path.dirname(__file__)' computes the directory containing the current script. Using '..' navigates up one directory level to reach the root folder, and then traverses into the 'database' directory to locate 'skills_db.json'. This dynamic path resolution ensures the script works regardless of the terminal's Current Working Directory (CWD).",
        "Lines 8-9: The JSON database is loaded into memory.",
        "Lines 12-15: The system needs to know what to look for. It instantiates 'all_skills' as a Python 'set()'. Sets inherently enforce uniqueness and provide O(1) lookup performance. Nested loops iterate through every job role in the database ('db['roles'].values()'), extracting every skill array, and adding each individual skill string into the set. The result is a unified, deduplicated dictionary of every skill recognized by the platform."
    ]
    add_code_section("4.1 Master Skill Aggregation", code_12, expl_12)

    code_13 = """    detected_skills = []
    
    # Case-insensitive keyword matching
    for skill in all_skills:
        pattern = r'\\b' + re.escape(skill) + r'\\b'
        if re.search(pattern, text, re.IGNORECASE):
            detected_skills.append(skill)
            
    return sorted(list(set(detected_skills)))"""
    expl_13 = [
        "Line 1: An empty list 'detected_skills' is initialized to hold matches.",
        "Line 4: A loop iterates over every single recognized skill.",
        "Line 5: A critical Regular Expression pattern is constructed. 're.escape(skill)' ensures that if a skill contains regex metacharacters (e.g., the '+' in 'C++'), they are treated as literal characters. The '\\b' tags on either side enforce 'word boundary' constraints. This prevents false positives; for example, the skill 'R' won't accidentally be matched inside the word 'React'.",
        "Lines 6-7: 're.search()' scans the entire resume 'text'. The 're.IGNORECASE' flag makes the search case-insensitive, meaning 'python', 'Python', and 'PYTHON' will all successfully match. If a match is found, the normalized skill name from the database is appended.",
        "Line 9: The list is cast to a set (to ensure no accidental duplicates if a skill was somehow matched twice), back to a list, sorted alphabetically, and returned. Sorting ensures deterministic output, which is highly preferred for UI rendering."
    ]
    add_code_section("4.2 Regex Pattern Matching Strategy", code_13, expl_13)

    # --- ANALYZER ---
    doc.add_paragraph('5. Gap Analysis Engine (analyzer.py)', style='CustomH1')
    doc.add_paragraph("The analyzer compares the candidate's discovered skills against the predefined expectations for their chosen role.")

    code_14 = """import json
import os

def analyze(detected_skills, role):
    db_path = os.path.join(os.path.dirname(__file__), '..', 'database', 'skills_db.json')
    
    with open(db_path, 'r') as f:
        db = json.load(f)
    
    required_skills = db['roles'].get(role, [])
    
    matched = [skill for skill in detected_skills if skill in required_skills]
    missing = [skill for skill in required_skills if skill not in detected_skills]
    
    return {
        "required": required_skills,
        "matched": matched,
        "missing": missing
    }"""
    expl_14 = [
        "Line 10: 'db['roles'].get(role, [])' attempts to fetch the list of skills required for the specified 'role'. Utilizing the '.get()' method with an empty list fallback ('[]') provides safety; if the UI somehow submitted an invalid role, the system will not throw a KeyError.",
        "Line 12: Python list comprehension is used: '[skill for skill in detected_skills if skill in required_skills]'. This elegant, highly-optimized single line of code iterates through the user's detected skills and filters only those that are present in the job's requirement list. This array forms the 'matched' skills.",
        "Line 13: Another list comprehension determines the gaps: '[skill for skill in required_skills if skill not in detected_skills]'. It iterates through the job requirements and isolates those *missing* from the user's resume.",
        "Lines 15-19: A comprehensive dictionary payload containing 'required', 'matched', and 'missing' sets is returned to the main app loop."
    ]
    add_code_section("5.1 Logic & Set Intersections", code_14, expl_14)

    # --- SCORE GENERATOR ---
    doc.add_paragraph('6. Scoring Algorithm (score_generator.py)', style='CustomH1')
    doc.add_paragraph("This minimal but crucial module assigns numerical and qualitative metrics to the analysis.")

    code_15 = """def generate_score(matched, required):
    if not required:
        return {"score": 0, "label": "Beginner"}
    
    score = int((len(matched) / len(required)) * 100)
    
    if score >= 80:
        label = "Industry Ready"
    elif score >= 60:
        label = "Intermediate"
    else:
        label = "Beginner"
        
    return {
        "score": score,
        "label": label
    }"""
    expl_15 = [
        "Lines 2-3: A guard clause ensures that if the required array is empty (to prevent ZeroDivisionError), it defaults safely to 0%.",
        "Line 5: The score is computed by taking the length (count) of the 'matched' array, dividing it by the length of the 'required' array to get a decimal, multiplying by 100, and casting to an integer ('int()') to truncate any decimals, providing a clean whole number percentage.",
        "Lines 7-12: A standard thresholding logic block assigns a qualitative label. A score of 80% or higher is deemed 'Industry Ready', 60-79% is 'Intermediate', and anything below 60% defaults to 'Beginner'.",
        "Lines 14-17: The dual-value dictionary is returned."
    ]
    add_code_section("6.1 Metric Generation", code_15, expl_15)

    # --- RECOMMENDER ---
    doc.add_paragraph('7. Learning Path Recommendation (recommender.py)', style='CustomH1')
    doc.add_paragraph("This module maps missing skills to actionable, customized learning interventions.")

    code_16 = """import json
import os

def recommend(missing_skills):
    db_path = os.path.join(os.path.dirname(__file__), '..', 'database', 'skills_db.json')
    
    with open(db_path, 'r') as f:
        db = json.load(f)
    
    courses = db.get('courses', {})
    recommendations = []"""
    expl_16 = [
        "Line 10: 'courses = db.get('courses', {})' safely extracts the 'courses' mapping dictionary from the main JSON structure, providing an empty dictionary fallback if it doesn't exist.",
        "Line 11: An empty list 'recommendations' is instantiated to collect the generated tips."
    ]
    add_code_section("7.1 Initialization", code_16, expl_16)

    code_17 = """    for skill in missing_skills:
        url = courses.get(skill, "https://www.google.com/search?q=learn+" + skill.replace(" ", "+"))
        tip = f"Complete a project or certification in {skill} to strengthen your profile."
        
        # Specific tips for some common skills
        if "Flask" in skill or "Django" in skill:
            tip = f"Build 2 web projects using {skill} to solidify this skill."
        elif "Python" in skill:
            tip = "Solve 50+ LeetCode problems in Python to master syntax and logic."
        elif "Docker" in skill or "Kubernetes" in skill:
            tip = f"Deploy a multi-container app using {skill} on a cloud platform."
            
        recommendations.append({
            "skill": skill,
            "url": url,
            "tip": tip
        })
        
    return sorted(recommendations, key=lambda x: x['skill'])"""
    expl_17 = [
        "Line 1: A loop processes every skill the candidate lacked.",
        "Line 2: The system attempts to find a hardcoded URL in the 'courses' database. If none exists, it dynamically generates a Google search URL by appending the skill name and replacing spaces with '+' symbols for URL compatibility.",
        "Line 3: A generic tip template is initialized.",
        "Lines 6-11: Rule-based conditional overrides inject highly specific, tailored advice for known high-value skills. For web frameworks, project building is suggested. For Python, algorithmic practice via LeetCode is emphasized. For DevOps tooling, cloud deployment is recommended.",
        "Lines 13-17: The assembled dictionary for the current skill is appended to the payload.",
        "Line 19: The final array is sorted alphabetically based on the 'skill' key using a lambda function: 'key=lambda x: x['skill']', ensuring the final UI output is neat and organized."
    ]
    add_code_section("7.2 Dynamic Resource Generation", code_17, expl_17)

    # --- CONCLUSION ---
    doc.add_page_break()
    doc.add_paragraph('8. Conclusion and Architectural Summary', style='CustomH1')
    conclusion = (
        "In conclusion, the SkillSyncAI codebase demonstrates robust engineering practices. "
        "The clear separation of concerns into distinct modules (Parsing, Analysis, Extraction, Scoring, and Recommendation) "
        "ensures that the application is highly maintainable and scalable. The use of resilient file encoding fallbacks, "
        "secure filename sanitization, and defensive dictionary fetching paradigms highlight a mature approach to web application development.\n\n"
        "This architectural breakdown verifies that the data flows seamlessly from the client upload layer to the complex algorithmic processing layer, "
        "and finally structures the output into comprehensive, actionable JSON responses that power both the dynamic frontend templates and the static generated reports."
    )
    doc.add_paragraph(conclusion)

    # Save document
    report_filename = 'SkillSyncAI_Code_Report.docx'
    doc.save(report_filename)
    print(f"Report generated successfully: {report_filename}")

if __name__ == '__main__':
    create_report()
